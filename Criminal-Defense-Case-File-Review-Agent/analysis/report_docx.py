"""把 deep_read_agentic.py / batch_analysis.py 产出的主报告(带简单 markdown 标记的纯文本)
渲染成 Word 文档,给律师当开庭前自己看的参考资料用——不是庭审出示文件,所以不需要剥离健康
检查提示之类的工作痕迹,报告里本来是什么内容就原样呈现,只是换个更好读的排版。

之前每次交付报告 docx 都是手工转一遍,案子一多容易漏、也不一致。报告本身只有几种固定的
markdown 标记(#/##标题、>引用说明、-列表、**加粗**、【引用】页码标记),不需要完整的
markdown 解析器,按行处理这几种情况就够。"""

import re

from docx import Document
from docx.shared import Pt, RGBColor

_CITATION_RE = re.compile(r"(【[^】]+】)")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

_CITATION_COLOR = RGBColor(0x22, 0x55, 0xAA)


def _add_inline_runs(paragraph, text: str) -> None:
    """处理一行里的 **加粗** 和 【引用】标记,两者可能同时出现(比如整段加粗里包一个引用)。"""
    pos = 0
    for m in _BOLD_RE.finditer(text):
        if m.start() > pos:
            _add_citation_runs(paragraph, text[pos : m.start()], bold=False)
        _add_citation_runs(paragraph, m.group(1), bold=True)
        pos = m.end()
    if pos < len(text):
        _add_citation_runs(paragraph, text[pos:], bold=False)


def _add_citation_runs(paragraph, text: str, bold: bool) -> None:
    last_index = 0
    for m in _CITATION_RE.finditer(text):
        if m.start() > last_index:
            run = paragraph.add_run(text[last_index : m.start()])
            run.bold = bold
            run.font.size = Pt(10.5)
        run = paragraph.add_run(m.group(0))
        run.bold = True
        run.font.color.rgb = _CITATION_COLOR
        run.font.size = Pt(10.5)
        last_index = m.end()
    if last_index < len(text):
        run = paragraph.add_run(text[last_index:])
        run.bold = bold
        run.font.size = Pt(10.5)


def build_report_docx(report_text: str, case_name: str, out_path: str) -> None:
    doc = Document()
    title = doc.add_heading(f"{case_name} —— 阅卷分析报告", level=0)
    title.alignment = 1  # 居中

    for raw_line in report_text.split("\n"):
        line = raw_line.rstrip()

        if not line:
            continue
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            _add_inline_runs(p, line)

    doc.save(out_path)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="把主报告渲染成 Word 文档")
    parser.add_argument("report_path", help="report.txt 路径")
    parser.add_argument("case_name", help="案子名称,用作文档标题")
    parser.add_argument("--out", default="阅卷分析报告.docx")
    args = parser.parse_args()

    with open(args.report_path, "r", encoding="utf-8") as f:
        text = f.read()

    build_report_docx(text, args.case_name, args.out)
    print(f"写入 {args.out}")
