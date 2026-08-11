"""把 indictment_check.py 产出的起诉书核对文本,渲染成 Word 表格。

格式选择的由来:最早试过纯文本(起诉书行用红底区分),用户反馈"光靠颜色不够,得看内容才能
判断哪几行是同一组对照"——改成表格,左栏起诉书、右栏证据,一行是一组对照,靠表格线天然
分组,不用再靠读内容猜边界。

原来是用 Node.js(docx 这个 npm 包)写的临时脚本,这里用 python-docx 重写成项目正式模块——
跟项目其余部分保持同一个技术栈,不需要额外装 Node 环境才能跑完整流程。"""

import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from analysis.indictment_check import parse_comparison_groups

_CITATION_RE = re.compile(r"(【[^】]+】)")

_INDICTMENT_COLOR = RGBColor(0xAA, 0x33, 0x22)
_EVIDENCE_COLOR = RGBColor(0x22, 0x55, 0xAA)


def _add_citation_runs(paragraph, text: str, color: RGBColor) -> None:
    last_index = 0
    for m in _CITATION_RE.finditer(text):
        if m.start() > last_index:
            paragraph.add_run(text[last_index : m.start()]).font.size = Pt(10)
        run = paragraph.add_run(m.group(0))
        run.bold = True
        run.font.color.rgb = color
        run.font.size = Pt(10)
        last_index = m.end()
    if last_index < len(text):
        paragraph.add_run(text[last_index:]).font.size = Pt(10)


def _set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    # python-docx 的 cell.width 在某些 Word 版本里需要同时设置底层 XML 才会真正生效,
    # 只设 cell.width 属性偶尔会被表格自动布局忽略。
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = tc_pr.makeelement(qn("w:tcW"), {})
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))  # 1cm ≈ 567 twips
    tc_w.set(qn("w:type"), "dxa")


def build_indictment_docx(comparison_text: str, case_name: str, out_path: str, doc_type: str = "起诉书") -> int:
    """doc_type 是"起诉书"或"起诉意见书"——决定标题、说明文字、表头怎么称呼这份文书,
    以及按哪个标记去分组(必须跟生成 comparison_text 时用的 doc_type 一致)。返回对照组数量。"""
    groups = parse_comparison_groups(comparison_text, doc_type=doc_type)

    doc = Document()
    title = doc.add_heading(f"{case_name} —— {doc_type}事实核对", level=0)
    title.alignment = 1  # 居中

    note = doc.add_paragraph()
    note_run = note.add_run(
        f"说明:本表只做并列对照,不判断{doc_type}与卷宗证据是否一致,不下任何结论——"
        f"左栏是{doc_type}怎么说,右栏是卷宗证据怎么说,一行是一组对照,判断权在律师。"
    )
    note_run.italic = True
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    indictment_header = f"{doc_type}表述"
    header_cells = table.rows[0].cells
    for cell, text in zip(header_cells, [indictment_header, "卷宗证据情况"]):
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        _set_cell_width(cell, 8.0 if text == indictment_header else 9.0)

    for g in groups:
        row_cells = table.add_row().cells
        indictment_cell, evidence_cell = row_cells
        _set_cell_width(indictment_cell, 8.0)
        _set_cell_width(evidence_cell, 9.0)

        indictment_cell.paragraphs[0].text = ""
        for i, line in enumerate(g["indictment"]):
            p = indictment_cell.paragraphs[0] if i == 0 else indictment_cell.add_paragraph()
            _add_citation_runs(p, line, _INDICTMENT_COLOR)

        evidence_cell.paragraphs[0].text = ""
        evidence_lines = g["evidence"] or ["(未找到对应证据)"]
        for i, line in enumerate(evidence_lines):
            p = evidence_cell.paragraphs[0] if i == 0 else evidence_cell.add_paragraph()
            if g["evidence"]:
                _add_citation_runs(p, line, _EVIDENCE_COLOR)
            else:
                run = p.add_run(line)
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(out_path)
    return len(groups)


if __name__ == "__main__":
    import argparse

    parser = argparse.ArgumentParser(description="把起诉书核对文本渲染成 Word 表格")
    parser.add_argument("comparison_path", help="indictment_check.py 产出的核对文本文件")
    parser.add_argument("case_name", help="案子名称,用作文档标题")
    parser.add_argument("--doc-type", default="起诉书", choices=["起诉书", "起诉意见书"])
    parser.add_argument("--out", default="起诉书事实核对.docx")
    args = parser.parse_args()

    with open(args.comparison_path, "r", encoding="utf-8") as f:
        text = f.read()

    n = build_indictment_docx(text, args.case_name, args.out, doc_type=args.doc_type)
    print(f"{n}组对照,写入 {args.out}")
